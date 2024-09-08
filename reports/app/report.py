import logging
import os
import subprocess

import msoffcrypto
from docx import Document
from PyPDF2 import PdfReader, PdfWriter

from .models import input_data
from .models.reports import meeting_transcript, official_report, unofficial_report


def generate_report(
    input_data: input_data.Meeting_Transcript_Data | input_data.Official_Protocol_Data | input_data.Unofficial_Protocol_Data,
) -> str:
    """Generate a report meeting.

    This function takes an input data object and generates a report based on the type of the data.
    The report is generated in the specified document type and with the specified password.

    Args:
        input_data (
        input_data.Meeting_Transcript_Data |
        input_data.Official_Protocol_Data |
        input_data.Unofficial_Protocol_Data
        ): Input data for generating the report.

    Returns:
        str: The path to the generated report file.
    """
    match type(input_data.data):
        case meeting_transcript.MeetingTranscript:
            # Generate a report for a meeting transcript
            report = generate_report_meeting(
                input_data.name_report, input_data.data, input_data.document_type, input_data.password
            )
            return report
        case official_report.OfficialProtocol:
            # Generate a report for an official protocol
            report = generate_official_report(
                input_data.name_report, input_data.data, input_data.document_type, input_data.password
            )
            return report
        case unofficial_report.UnofficialProtocol:
            # Generate a report for an unofficial protocol
            report = generate_unofficial_report(
                input_data.name_report, input_data.data, input_data.document_type, input_data.password
            )
            return report
        case _:
            # Raise an error if the type of the input data is unknown
            raise ValueError("Unknown data type")


def generate_official_report(
    name_report: str, data: official_report.OfficialProtocol, document_type: input_data.Type_Document, password: str | None = None
) -> str:
    """Generate an official protocol report.

    Args:
        name_report: The name of the report.
        data: The official protocol data.
        document_type: The type of document (docx or pdf).
        password: The password for encrypted documents (optional).

    Returns:
        str: The path to the generated report file.
    """
    document = generate_official_report_docx(name_report, data)

    if document_type == input_data.Type_Document.docx and password == "":
        return document
    if document_type == input_data.Type_Document.pdf:
        return generate_pdf(document, name_report, password)
    elif password != "":
        return encrypt_docx_report(document, name_report, password)


def generate_official_report_docx(name_report: str, data: official_report.OfficialProtocol) -> str:
    """Generate an official protocol report in docx format.

    Args:
        name_report: The name of the report.
        data: The official protocol data.

    Returns:
        str: The path to the generated report file.
    """
    document = get_document()

    document.add_heading(name_report, level=1)
    if data.time is not None:
        document.add_paragraph(f"Время: {data.time.strftime('%H:%M:%S')}")
    if data.date is not None:
        document.add_paragraph(f"Дата проведения: {data.date.strftime('%d-%m-%Y')}")

    if data.attendees is not None:
        document.add_paragraph("Участники:\n" + "\n".join(f"{value}" for value in data.attendees)).style = "List Bullet"

    for block in data.blocks:
        document.add_heading(block.name_block, level=2)

        for proposal in block.proposals:
            if proposal.text != "":
                document.add_heading(proposal.text, level=3)
            if proposal.context != "":
                document.add_paragraph(proposal.context)
            if proposal.audio_time is not None:
                document.add_paragraph(
                    f"Время в аудиозаписи:\n\tc {proposal.audio_time.start.strftime('%H:%M:%S')}"
                    + f" до {proposal.audio_time.end.strftime('%H:%M:%S')}"
                )

    document.add_page_break()

    for errand in data.errand_protocol.list_errands:
        if errand.assignee != "":
            document.add_heading(errand.assignee, level=3)
        if errand.deadline is not None:
            document.add_heading(f"Дедлайн: {errand.deadline}", level=4)
        if errand.context != "":
            document.add_paragraph(errand.context)

    document.save(in_reports(name_report + ".docx"))
    return in_reports(name_report + ".docx")


def in_reports(name_report: str) -> str:
    """Join the reports directory with the name of the report.

    Args:
        name_report (str): The name of the report.

    Returns:
        str: The full path to the report.
    """
    return os.path.join("reports", name_report)


def generate_pdf(document: str, name_report: str, password: str) -> str:
    """Generate a PDF report.

    Args:
        document (str): The path to the document to convert to PDF.
        name_report (str): The name of the report.
        password (str): The password to encrypt the report with.

    Returns:
        str: The path to the generated PDF report.
    """
    try:
        subprocess.run(["/usr/bin/abiword", "--to=pdf", document], check=False)
    except subprocess.CalledProcessError as e:
        logging.error(e)
    if password != "":
        with (
            open(in_reports(f"{name_report}.pdf"), "rb") as file,
            open(in_reports(f"encrypt_{name_report}.pdf"), "wb") as encrypt_report,
        ):
            pdf_reader = PdfReader(file)
            pdf_writer = PdfWriter()
            for page in pdf_reader.pages:
                pdf_writer.add_page(page)
            pdf_writer.encrypt(password)
            pdf_writer.write(encrypt_report)
            return encrypt_report.name
    return in_reports(name_report + ".pdf")


def encrypt_docx_report(
    document: str,  # The path to the document to encrypt.
    name_report: str,  # The name of the report.
    password: str,  # The password to encrypt the report with.
) -> str:  # The path to the generated encrypted report.
    """Encrypt a docx report.

    The report is encrypted with the given password and saved to a file with
    the same name as the input document, but with "encrypt_" prefixed to the
    name.

    Args:
        document: The path to the document to encrypt.
        name_report: The name of the report.
        password: The password to encrypt the report with.

    Returns:
        str: The path to the generated encrypted report.
    """
    with open(document, "rb") as file, open(in_reports(f"encrypt_{name_report}.docx"), "wb") as encrypt_report:
        office_file = msoffcrypto.OfficeFile(file)
        office_file.load_key(password)
        office_file.encrypt(password, encrypt_report)
        return encrypt_report.name


def generate_unofficial_report(
    name_report: str,
    data: unofficial_report.UnofficialProtocol,
    document_type: input_data.Type_Document,
    password: str | None = None,
) -> str:
    """Generate an unofficial report meeting.

    Args:
        name_report: The name of the report.
        data: The unofficial protocol data.
        document_type: The type of document (docx or pdf).
        password: The password for encrypted documents (optional).

    Returns:
        str: The path to the generated report file.
    """
    document = generate_unofficial_report_docx(name_report, data)

    if document_type == input_data.Type_Document.docx and password == "":
        return document
    if document_type == input_data.Type_Document.pdf:
        return generate_pdf(document, name_report, password)
    elif password != "":
        return encrypt_docx_report(document, name_report, password)


def generate_unofficial_report_docx(name_report: str, data: unofficial_report.UnofficialProtocol) -> str:
    """Generate an unofficial report meeting in docx format.

    Args:
        name_report: The name of the report.
        data: The unofficial protocol data.

    Returns:
        str: The path to the generated report file.
    """
    document = get_document()

    document.add_heading(name_report, level=1)
    if data.time is not None:
        document.add_paragraph(f"Время: {data.time.strftime('%H:%M:%S')}")
    if data.date is not None:
        document.add_paragraph(f"Дата проведения: {data.date.strftime('%d-%m-%Y')}")
    if data.duration is not None:
        document.add_paragraph(f"Длительность: {data.duration}")

    if data.participants is not None:
        document.add_paragraph("Участники:\n" + "\n".join(value for value in data.participants)).style = "List Bullet"

    if data.agenda is not None:
        document.add_paragraph("Повестка дня:\n" + "\n".join(value for value in data.agenda)).style = "List Bullet"

    for block in data.blocks:
        document.add_heading(block.name_block, level=2)

        for proposal in block.proposals:
            if proposal.text != "":
                document.add_heading(proposal.text, level=3)
            if proposal.context != "":
                document.add_paragraph(proposal.context)
            if proposal.audio_time is not None:
                document.add_paragraph(
                    f"Время в аудиозаписи: {proposal.audio_time.start.strftime('%H:%M:%S')}"
                    + f" - {proposal.audio_time.end.strftime('%H:%M:%S')}"
                )

    document.save(in_reports(name_report + ".docx"))
    return in_reports(name_report + ".docx")


def generate_report_meeting(
    name_report: str,
    data: meeting_transcript.MeetingTranscript,
    document_type: input_data.Type_Document,
    password: str | None = None,
) -> str:
    """Generate a meeting transcript report.

    Args:
        name_report: The name of the report.
        data: The meeting transcript data.
        document_type: The type of document (docx or pdf).
        password: The password for encrypted documents (optional).

    Returns:
        str: The path to the generated report file.
    """
    document = generate_meeting_report_docx(name_report, data)

    if document_type == input_data.Type_Document.docx and password == "":
        return document
    if document_type == input_data.Type_Document.pdf:
        return generate_pdf(document, name_report, password)
    elif password != "":
        return encrypt_docx_report(document, name_report, password)


def generate_meeting_report_docx(name_report: str, data: meeting_transcript.MeetingTranscript) -> str:
    """Generate a meeting transcript report in docx format.

    Args:
        name_report: The name of the report.
        data: The meeting transcript data.

    Returns:
        str: The path to the generated report file.
    """
    document = get_document()

    document.add_heading(name_report, level=1)

    for speaker in data.speakers_transcript:
        if speaker.speaker_name != "":
            document.add_heading(f"ФИО: {speaker.speaker_name}", level=2)
        if speaker.audio_time is not None:
            document.add_heading("Время в аудиозаписи:", level=3)
            document.add_paragraph(
                f"Старт: {speaker.audio_time.start.strftime('%H:%M:%S')}" +
                f"\n Конец: {speaker.audio_time.end.strftime('%H:%M:%S')}"
            )
        if speaker.transcript_text != "":
            document.add_paragraph(speaker.transcript_text)
        document.add_page_break()

    name_report_docx = f"{name_report}.docx"
    document.save(in_reports(name_report_docx))

    return in_reports(name_report_docx)


def get_document() -> Document:
    """Create a new document format docx.

    Returns:
        Document: A new document.
    """
    document = Document()
    return document
